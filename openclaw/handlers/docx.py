#!/usr/bin/env python3
"""Word Document Handler (.docx)"""

from pathlib import Path
from typing import Dict, Any, List
from ..core.document import Document, DocumentHandler


class DocxHandler(DocumentHandler):
    """Handler for Microsoft Word documents (.docx).
    
    Uses python-docx library for reading and writing.
    """
    
    def get_supported_formats(self) -> List[str]:
        return ["docx", "word"]
    
    def supports_format(self, format_name: str) -> bool:
        return format_name.lower() in ["docx", "word"]
    
    def read(self, source: str) -> Document:
        """Read Word document.
        
        Args:
            source: File path to .docx file
            
        Returns:
            Document object
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ImportError: If python-docx not installed
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        path = Path(source)
        
        if not path.exists():
            raise FileNotFoundError(f"Word document not found: {source}")
        
        # Read document
        docx_doc = DocxDocument(str(path))
        
        # Extract text from paragraphs
        paragraphs = []
        for para in docx_doc.paragraphs:
            paragraphs.append(para.text)
        
        # Extract text from tables
        for table in docx_doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(row_text))
        
        content = "\n\n".join(paragraphs)
        
        # Gather metadata
        stat = path.stat()
        metadata = {
            "path": str(path.absolute()),
            "size": stat.st_size,
            "format": "docx",
            "paragraph_count": len(docx_doc.paragraphs),
            "table_count": len(docx_doc.tables),
            "sections": len(docx_doc.sections)
        }
        
        # Extract core properties if available
        try:
            core_props = docx_doc.core_properties
            metadata.update({
                "title": core_props.title,
                "author": core_props.author,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
            })
        except Exception:
            pass
        
        return Document(
            content=content,
            metadata=metadata,
            format="docx",
            source=source
        )
    
    def write(self, destination: str, document: Document) -> Dict[str, Any]:
        """Write Word document.
        
        Args:
            destination: File path to write
            document: Document object
            
        Returns:
            Result dictionary
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        path = Path(destination)
        
        # Create parent directories if needed
        path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create new document
        docx_doc = DocxDocument()
        
        # Add content - split by paragraphs
        content = document.content
        paragraphs = content.split("\n\n")
        
        for para_text in paragraphs:
            if para_text.strip():
                docx_doc.add_paragraph(para_text.strip())
        
        # Save document
        docx_doc.save(str(path))
        
        return {
            "success": True,
            "data": {
                "path": str(path.absolute()),
                "size": path.stat().st_size,
                "written": True,
                "format": "docx"
            }
        }
